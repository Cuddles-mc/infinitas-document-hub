"""Shortlist Generator — upload CVs, review extracted data, download branded PPTX."""

import io
import streamlit as st
from ui import page_header, step_flow, form_section


QUAL_KEYWORDS = ["chartered", "ca ", "ca,", "cpa", "cima", "cfa", "iod", "member", "fellow", "certified", "registered", "accredited"]


def _build_diff_html(original: str, corrected: str) -> str:
    """Build HTML showing word-level differences between original and corrected text.

    Deletions shown in red strikethrough, insertions in green bold.
    Returns empty string if texts are identical.
    """
    import difflib

    if original.strip() == corrected.strip():
        return ""

    orig_words = original.split()
    corr_words = corrected.split()
    sm = difflib.SequenceMatcher(None, orig_words, corr_words)

    parts = []
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            parts.append(" ".join(orig_words[i1:i2]))
        elif op == "replace":
            old = " ".join(orig_words[i1:i2])
            new = " ".join(corr_words[j1:j2])
            parts.append(
                f'<span style="background:#fee2e2;color:#991b1b;text-decoration:line-through;">{old}</span> '
                f'<span style="background:#dcfce7;color:#166534;font-weight:600;">{new}</span>'
            )
        elif op == "delete":
            old = " ".join(orig_words[i1:i2])
            parts.append(
                f'<span style="background:#fee2e2;color:#991b1b;text-decoration:line-through;">{old}</span>'
            )
        elif op == "insert":
            new = " ".join(corr_words[j1:j2])
            parts.append(
                f'<span style="background:#dcfce7;color:#166534;font-weight:600;">{new}</span>'
            )

    return f'<div style="line-height:1.8;font-size:0.95rem;padding:0.75rem;border:1px solid #e5e7eb;border-radius:0.5rem;background:#fafafa;">{" ".join(parts)}</div>'


def _split_edu_qual(combined: str) -> tuple[str, str]:
    """Split a combined education/qualifications string into separate fields.

    Lines containing professional qualification keywords go to quals,
    everything else goes to education.
    """
    if not combined:
        return ("", "")

    edu_parts = []
    qual_parts = []

    for line in combined.replace("\x0b", "\n").split("\n"):
        line = line.strip()
        if not line:
            continue
        if any(kw in line.lower() for kw in QUAL_KEYWORDS):
            qual_parts.append(line)
        else:
            edu_parts.append(line)

    return (", ".join(edu_parts), ", ".join(qual_parts))


def render():
    page_header("Shortlist Generator", "Upload CVs to create a branded shortlist presentation")

    # Determine current step
    has_candidates = "sl_candidates" in st.session_state
    has_pptx = "sl_pptx_bytes" in st.session_state
    current_step = 2 if has_pptx else (1 if has_candidates else 0)
    step_flow(["Upload CVs", "Review & Edit", "Download"], current_step)

    if has_pptx:
        _render_download()
    elif has_candidates:
        _render_review()
    else:
        _render_upload()


# ---------------------------------------------------------------------------
# Step 1: Upload
# ---------------------------------------------------------------------------
def _render_upload():
    mode = st.radio(
        "What would you like to do?",
        ["Create new shortlist", "Add to existing shortlist"],
        key="sl_mode",
        horizontal=True,
    )
    st.markdown("")

    if mode == "Add to existing shortlist":
        _render_upload_append()
    else:
        _render_upload_new()


def _render_upload_new():
    form_section("Assignment Details")
    st.markdown(
        """
        <div style="background:#fef3c7;border-left:4px solid #f59e0b;padding:0.75rem 1rem;
                    border-radius:0.375rem;margin-bottom:0.5rem;font-size:0.9rem;color:#78350f;">
          <strong>⚠ Double-check the client name.</strong> It appears on the shortlist title slide
          and — if you generate CV Profiles afterwards — on <em>every</em> cover page and output
          filename. Spelling, capitalisation, and "Limited" vs "Ltd" all flow through to the client.
        </div>
        """,
        unsafe_allow_html=True,
    )
    col1, col2 = st.columns(2)
    with col1:
        client_name = st.text_input("Client name *", key="sl_client_name_input")
    with col2:
        role_title = st.text_input("Role title *", key="sl_role_title_input")

    form_section("Template")
    template_choice = st.radio(
        "Choose template",
        ["Standard (one page, portrait)", "Executive (two page, landscape)"],
        key="sl_template_choice",
        horizontal=True,
        help="Standard works for most roles. Executive gives more space for detailed notes on senior appointments.",
    )

    if template_choice.startswith("Executive"):
        col_prep, col_date = st.columns(2)
        with col_prep:
            prepared_by = st.text_input(
                "Prepared by",
                value=st.session_state.get("ms_user", ""),
                key="sl_prepared_by_input",
            )
        with col_date:
            from datetime import datetime
            prepared_date = st.text_input(
                "Date",
                value=datetime.now().strftime("%B %Y"),
                key="sl_prepared_date_input",
            )

    form_section("Upload CVs")
    uploaded_files = st.file_uploader(
        "Upload candidate CVs (PDF or DOCX)",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key="sl_cv_upload",
    )

    if uploaded_files:
        st.caption(f"{len(uploaded_files)} file(s) selected")

    st.markdown("")
    col_btn, _ = st.columns([1, 2])
    with col_btn:
        extract = st.button(
            "Extract candidate data",
            type="primary",
            width="stretch",
            disabled=not uploaded_files,
        )

    if extract:
        if not client_name or not role_title:
            st.error("Please fill in client name and role title.")
            return
        if not uploaded_files:
            st.error("Please upload at least one CV.")
            return

        candidates = []
        progress = st.progress(0, text="Extracting CVs...")

        for i, f in enumerate(uploaded_files):
            progress.progress(
                (i) / len(uploaded_files),
                text=f"Extracting {f.name}...",
            )

            try:
                # Store raw bytes before reading for text extraction
                cv_raw = f.read()
                f.seek(0)

                cv_text = _extract_text_from_bytes(cv_raw, f.name)
                if not cv_text.strip():
                    st.warning(f"Could not extract text from {f.name}")
                    continue

                from ai import extract_cv_data
                data = extract_cv_data(cv_text)

                # Add defaults and store original CV bytes
                for entry in data.get("career", []):
                    entry["include"] = True
                data["source_file"] = f.name
                data["cv_bytes"] = cv_raw
                data["notes"] = ""
                data["use_lorem"] = True
                candidates.append(data)

            except Exception as e:
                st.error(f"Error processing {f.name}: {e}")

        progress.progress(1.0, text="Done!")

        if candidates:
            st.session_state.sl_candidates = candidates
            st.session_state.sl_client_name = client_name
            st.session_state.sl_role_title = role_title
            st.session_state.sl_template = "executive" if template_choice.startswith("Executive") else "standard"
            if template_choice.startswith("Executive"):
                st.session_state.sl_prepared_by = prepared_by
                st.session_state.sl_prepared_date = prepared_date
            st.rerun()
        else:
            st.error("No candidates could be extracted.")


def _render_upload_append():
    """Upload flow for adding candidates to an existing shortlist."""
    form_section("Existing Shortlist")
    existing_file = st.file_uploader(
        "Upload your existing shortlist PPTX",
        type=["pptx"],
        key="sl_existing_pptx",
    )

    form_section("New CVs")
    uploaded_files = st.file_uploader(
        "Upload new candidate CVs (PDF or DOCX)",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key="sl_new_cv_upload",
    )

    if uploaded_files:
        st.caption(f"{len(uploaded_files)} file(s) selected")

    st.markdown("")
    col_btn, _ = st.columns([1, 2])
    with col_btn:
        extract = st.button(
            "Extract candidate data",
            type="primary",
            width="stretch",
            disabled=not uploaded_files or not existing_file,
        )

    if extract:
        existing_bytes = existing_file.read()

        # Detect template type from slide dimensions
        from pptx import Presentation as PptxPresentation
        prs = PptxPresentation(io.BytesIO(existing_bytes))
        is_landscape = prs.slide_width > prs.slide_height
        template_type = "executive" if is_landscape else "standard"

        candidates = []
        progress = st.progress(0, text="Extracting CVs...")

        for i, f in enumerate(uploaded_files):
            progress.progress(i / len(uploaded_files), text=f"Extracting {f.name}...")
            try:
                cv_raw = f.read()
                f.seek(0)
                cv_text = _extract_text_from_bytes(cv_raw, f.name)
                if not cv_text.strip():
                    st.warning(f"Could not extract text from {f.name}")
                    continue
                from ai import extract_cv_data
                data = extract_cv_data(cv_text)
                for entry in data.get("career", []):
                    entry["include"] = True
                data["source_file"] = f.name
                data["cv_bytes"] = cv_raw
                data["notes"] = ""
                data["use_lorem"] = True
                candidates.append(data)
            except Exception as e:
                st.error(f"Error processing {f.name}: {e}")

        progress.progress(1.0, text="Done!")

        if candidates:
            st.session_state.sl_candidates = candidates
            st.session_state.sl_existing_pptx = existing_bytes
            st.session_state.sl_existing_filename = existing_file.name
            st.session_state.sl_template = template_type
            st.session_state.sl_append_mode = True
            st.session_state.sl_client_name = ""
            st.session_state.sl_role_title = ""
            st.rerun()
        else:
            st.error("No candidates could be extracted.")


def _extract_text(uploaded_file) -> str:
    """Extract plain text from an uploaded PDF or DOCX file."""
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    return _extract_text_from_bytes(data, name)


def _extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract plain text from raw bytes given the filename for format detection."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf_text(data)
    elif name.endswith(".docx"):
        return _extract_docx_text(data)
    return ""


def _extract_pdf_text(data: bytes) -> str:
    """Extract text from PDF bytes."""
    import io
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except ImportError:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("Install pypdf or PyPDF2: pip install pypdf")


def _extract_docx_text(data: bytes) -> str:
    """Extract text from DOCX bytes."""
    import io
    from docx import Document
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _parse_notes_docx(uploaded_file, candidates: list[dict]) -> dict[int, str]:
    """Parse a DOCX with candidate notes, match sections to candidates.

    Detects candidate name lines by matching against known candidate names,
    regardless of whether they're styled as headings or plain text.
    Returns {candidate_index: notes_text} for matched candidates.
    """
    import io
    from docx import Document

    data = uploaded_file.read()
    uploaded_file.seek(0)
    doc = Document(io.BytesIO(data))

    # Build candidate name lookup (normalised lowercase -> index)
    # Full names first, then last names (full names take priority)
    full_names = {}
    last_names = {}
    for i, cand in enumerate(candidates):
        name = cand.get("name", "").strip()
        if name:
            full_names[name.lower()] = i
            parts = name.split()
            if len(parts) > 1:
                last_names[parts[-1].lower()] = i

    def _match_name(text: str) -> int | None:
        """Check if a line matches a candidate name. Short lines only."""
        text_lower = text.lower().strip()
        # Name lines are short — skip anything that looks like a sentence
        if len(text_lower.split()) > 5:
            return None
        # Try full name match first
        for name_key, idx in full_names.items():
            if name_key in text_lower or text_lower in name_key:
                return idx
        # Fall back to last name
        for name_key, idx in last_names.items():
            if text_lower == name_key or text_lower.startswith(name_key) or text_lower.endswith(name_key):
                return idx
        return None

    # Walk paragraphs
    sections: dict[int, list[str]] = {}
    current_idx = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        matched = _match_name(text)
        if matched is not None:
            current_idx = matched
            sections[current_idx] = []
            continue

        if current_idx is not None:
            sections[current_idx].append(text)

    # Join each candidate's paragraphs with double newlines
    return {idx: "\n\n".join(lines) for idx, lines in sections.items() if lines}


# ---------------------------------------------------------------------------
# Step 2: Review & Edit
# ---------------------------------------------------------------------------
def _render_review():
    # Back button
    if st.button("< Back to upload"):
        for key in ("sl_candidates", "sl_existing_pptx", "sl_existing_filename", "sl_append_mode"):
            st.session_state.pop(key, None)
        st.rerun()

    client_name = st.session_state.get("sl_client_name", "")
    role_title = st.session_state.get("sl_role_title", "")
    template_label = "Executive (landscape)" if st.session_state.get("sl_template") == "executive" else "Standard"
    if st.session_state.get("sl_append_mode"):
        existing_name = st.session_state.get("sl_existing_filename", "existing shortlist")
        st.info(f"Adding {len(st.session_state.sl_candidates)} new candidate(s) to **{existing_name}** — {template_label}")
    else:
        st.info(f"**{role_title}** at **{client_name}** — {len(st.session_state.sl_candidates)} candidate(s) — {template_label}")

    candidates = st.session_state.sl_candidates

    # Upload notes document
    form_section("Import Notes")
    notes_file = st.file_uploader(
        "Upload a Word document with candidate notes (optional)",
        type=["docx"],
        key="sl_notes_upload",
        help="One document with candidate names as headings, notes below each.",
    )
    if notes_file:
        col_import, _ = st.columns([1, 2])
        with col_import:
            if st.button("Import notes", type="primary", key="sl_import_notes"):
                notes_map = _parse_notes_docx(notes_file, candidates)
                if notes_map:
                    for idx, notes_text in notes_map.items():
                        candidates[idx]["notes"] = notes_text
                        candidates[idx]["use_lorem"] = False
                    matched = [candidates[i]["name"] for i in notes_map]
                    st.success(f"Imported notes for: {', '.join(matched)}")
                    st.rerun()
                else:
                    st.warning("Could not match any candidate names in the document.")

    for idx, cand in enumerate(candidates):
        with st.expander(f"**{cand.get('name', 'Unknown')}** ({cand.get('source_file', '')})", expanded=True):
            _render_candidate_editor(idx, cand)

    # Add manual candidate
    st.divider()
    if st.button("+ Add candidate manually"):
        candidates.append({
            "name": "",
            "career": [],
            "education_qualifications": "",
            "notice_period": "",
            "salary_expectation": "",
            "notes": "",
            "use_lorem": True,
            "source_file": "Manual entry",
        })
        st.rerun()

    # Generate button
    st.markdown("")
    col_btn, _ = st.columns([1, 2])
    with col_btn:
        generate = st.button(
            "Generate Shortlist PPTX",
            type="primary",
            width="stretch",
        )

    if generate:
        # Validate
        valid_candidates = [c for c in candidates if c.get("name", "").strip()]
        if not valid_candidates:
            st.error("At least one candidate needs a name.")
            return

        with st.spinner("Generating shortlist PPTX..."):
            try:
                append_mode = st.session_state.get("sl_append_mode", False)
                use_executive = st.session_state.get("sl_template") == "executive"

                if append_mode:
                    existing_bytes = st.session_state.sl_existing_pptx
                    if use_executive:
                        from generators.shortlist_executive_pptx import append_candidates
                        pptx_bytes = append_candidates(existing_bytes, valid_candidates)
                    else:
                        from generators.shortlist_pptx import append_candidates
                        pptx_bytes = append_candidates(existing_bytes, valid_candidates)
                    filename = st.session_state.get("sl_existing_filename", "shortlist.pptx")
                else:
                    if use_executive:
                        from generators.shortlist_executive_pptx import generate_executive_shortlist
                        pptx_bytes = generate_executive_shortlist(
                            client_name=client_name,
                            role_title=role_title,
                            candidates=valid_candidates,
                            prepared_by=st.session_state.get("sl_prepared_by", ""),
                            prepared_date=st.session_state.get("sl_prepared_date", ""),
                        )
                    else:
                        from generators.shortlist_pptx import generate_shortlist
                        pptx_bytes = generate_shortlist(
                            client_name=client_name,
                            role_title=role_title,
                            candidates=valid_candidates,
                        )
                    filename = f"{role_title} Shortlist prepared for {client_name} by Infinitas.pptx"

                st.session_state.sl_pptx_bytes = pptx_bytes
                st.session_state.sl_pptx_filename = filename
            except Exception as e:
                st.error(f"Error generating PPTX: {e}")
                import traceback
                st.code(traceback.format_exc())
                return

        st.rerun()


def _render_candidate_editor(idx: int, cand: dict):
    """Render the editable form for a single candidate."""

    # Name
    cand["name"] = st.text_input(
        "Candidate name",
        value=cand.get("name", ""),
        key=f"cand_name_{idx}",
    )

    # Photo
    form_section("Candidate Photo")
    photo_file = st.file_uploader(
        "Upload photo (optional — placeholder used if empty)",
        type=["png", "jpg", "jpeg"],
        key=f"cand_photo_{idx}",
    )
    if photo_file:
        cand["photo"] = photo_file.read()
        st.image(cand["photo"], width=150)
    else:
        cand["photo"] = None
        st.caption("No photo uploaded — placeholder will be used. You can swap it in PowerPoint later.")

    # Career history with checkboxes
    form_section("Career History")

    career = cand.get("career", [])
    if career:
        for ci, entry in enumerate(career):
            cols = st.columns([0.5, 3, 3, 1.5, 1.5, 1.5])
            with cols[0]:
                entry["include"] = st.checkbox(
                    "Inc",
                    value=entry.get("include", True),
                    key=f"career_inc_{idx}_{ci}",
                    label_visibility="collapsed",
                )
            with cols[1]:
                entry["company"] = st.text_input(
                    "Company",
                    value=entry.get("company", ""),
                    key=f"career_co_{idx}_{ci}",
                    label_visibility="collapsed" if ci > 0 else "visible",
                )
            with cols[2]:
                entry["title"] = st.text_input(
                    "Title",
                    value=entry.get("title", ""),
                    key=f"career_title_{idx}_{ci}",
                    label_visibility="collapsed" if ci > 0 else "visible",
                )
            with cols[3]:
                entry["start_date"] = st.text_input(
                    "Start",
                    value=entry.get("start_date", ""),
                    key=f"career_start_{idx}_{ci}",
                    label_visibility="collapsed" if ci > 0 else "visible",
                )
            with cols[4]:
                entry["end_date"] = st.text_input(
                    "End",
                    value=entry.get("end_date", ""),
                    key=f"career_end_{idx}_{ci}",
                    label_visibility="collapsed" if ci > 0 else "visible",
                )
            with cols[5]:
                if ci == 0:
                    st.caption("Delete")
                if st.button("X", key=f"career_del_{idx}_{ci}"):
                    career.pop(ci)
                    st.rerun()
    else:
        st.caption("No career entries. Add one below.")

    if st.button("+ Add career entry", key=f"add_career_{idx}"):
        career.append({
            "company": "",
            "title": "",
            "start_date": "",
            "end_date": "Present",
            "include": True,
        })
        st.rerun()

    # Split combined field on first render only
    if "education" not in cand and "education_qualifications" in cand:
        edu_split, qual_split = _split_edu_qual(cand.get("education_qualifications", ""))
        cand["education"] = edu_split
        cand["professional_qualifications"] = qual_split
        cand["show_education"] = bool(edu_split.strip())
        cand["show_prof_quals"] = bool(qual_split.strip())

    # Details — mirrors the PPTX Details table row order
    form_section("Details")

    # Notice period
    cand["notice_period"] = st.text_input(
        "Notice period",
        value=cand.get("notice_period", ""),
        key=f"cand_notice_{idx}",
        placeholder="e.g. Available immediately, 4 weeks",
    )

    # Salary expectation
    cand["salary_expectation"] = st.text_input(
        "Salary expectation",
        value=cand.get("salary_expectation", ""),
        key=f"cand_salary_{idx}",
        placeholder="e.g. $250,000 - $280,000",
    )

    # Education (with Include toggle)
    col_edu, col_edu_check = st.columns([4, 1])
    with col_edu:
        cand["education"] = st.text_input(
            "Education",
            value=cand.get("education", ""),
            key=f"cand_edu_{idx}",
            placeholder="e.g. Bachelor of Commerce, University of Auckland",
        )
    with col_edu_check:
        st.markdown("<div style='height:1.85rem'></div>", unsafe_allow_html=True)
        cand["show_education"] = st.checkbox(
            "Include",
            value=cand.get("show_education", bool(cand.get("education", "").strip())),
            key=f"cand_show_edu_{idx}",
        )

    # Professional Qualifications (with Include toggle)
    col_qual, col_qual_check = st.columns([4, 1])
    with col_qual:
        cand["professional_qualifications"] = st.text_input(
            "Professional qualifications",
            value=cand.get("professional_qualifications", ""),
            key=f"cand_quals_{idx}",
            placeholder="e.g. Chartered Accountant (CA), CAANZ",
        )
    with col_qual_check:
        st.markdown("<div style='height:1.85rem'></div>", unsafe_allow_html=True)
        cand["show_prof_quals"] = st.checkbox(
            "Include",
            value=cand.get("show_prof_quals", bool(cand.get("professional_qualifications", "").strip())),
            key=f"cand_show_quals_{idx}",
        )

    # Notes
    form_section("Consultant Notes")
    cand["use_lorem"] = st.checkbox(
        "No notes — use placeholder (lorem ipsum)",
        value=cand.get("use_lorem", True),
        key=f"cand_lorem_{idx}",
    )

    if not cand["use_lorem"]:
        cand["notes"] = st.text_area(
            "Notes",
            value=cand.get("notes", ""),
            height=200,
            key=f"cand_notes_{idx}",
            label_visibility="collapsed",
            placeholder="Write your assessment of this candidate...",
        )

        # Proofread button
        col_pr, _ = st.columns([1, 2])
        with col_pr:
            if st.button("Proofread notes", key=f"proofread_{idx}"):
                notes_text = cand.get("notes", "")
                if notes_text.strip():
                    with st.spinner("Proofreading..."):
                        try:
                            from ai import proofread_notes
                            corrected = proofread_notes(notes_text)
                            st.session_state[f"proofread_result_{idx}"] = corrected
                        except Exception as e:
                            st.error(f"Proofreading failed: {e}")

        # Show proofread result with highlighted changes
        proofread_key = f"proofread_result_{idx}"
        if proofread_key in st.session_state:
            corrected = st.session_state[proofread_key]
            original = cand.get("notes", "")
            diff_html = _build_diff_html(original, corrected)
            if diff_html:
                st.markdown("**Changes found:**")
                st.markdown(diff_html, unsafe_allow_html=True)
            else:
                st.success("No changes needed.")
            col_accept, col_reject, _ = st.columns([1, 1, 2])
            with col_accept:
                if st.button("Accept changes", key=f"accept_proof_{idx}", type="primary"):
                    cand["notes"] = corrected
                    del st.session_state[proofread_key]
                    st.rerun()
            with col_reject:
                if st.button("Keep original", key=f"reject_proof_{idx}"):
                    del st.session_state[proofread_key]
                    st.rerun()


# ---------------------------------------------------------------------------
# Step 3: Download
# ---------------------------------------------------------------------------
def _render_download():
    # Guard: sl_pptx_bytes can survive when its companions (client/role/candidates)
    # are popped by Back navigation or cleared by Streamlit Cloud between sessions.
    # If that happens, drop the dangling PPTX state and drop back to upload rather
    # than crash on attribute access.
    required = ("sl_client_name", "sl_role_title", "sl_candidates", "sl_pptx_filename")
    if not all(k in st.session_state for k in required):
        for k in ("sl_pptx_bytes", "sl_pptx_filename"):
            st.session_state.pop(k, None)
        st.rerun()

    # Back button
    if st.button("< Back to review"):
        del st.session_state["sl_pptx_bytes"]
        del st.session_state["sl_pptx_filename"]
        st.rerun()

    client_name = st.session_state.get("sl_client_name", "")
    role_title = st.session_state.get("sl_role_title", "")
    candidates = st.session_state.sl_candidates
    filename = st.session_state.sl_pptx_filename

    st.success(f"Shortlist generated: **{role_title}** at **{client_name}**")

    # Summary
    form_section("Summary")
    for cand in candidates:
        included_roles = sum(1 for c in cand.get("career", []) if c.get("include", True))
        st.markdown(f"- **{cand.get('name', 'Unknown')}** — {included_roles} career entries")

    # Shortlist PPTX download
    form_section("Shortlist Presentation")
    pptx_mime = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    st.download_button(
        label="Download Shortlist (.pptx)",
        data=st.session_state.sl_pptx_bytes,
        file_name=filename,
        mime=pptx_mime,
        type="primary",
        key="dl_pptx",
    )
    st.caption("Fully editable — adjust fonts, layout, and content in PowerPoint.")

    # CV Profiles — reuse already-uploaded CVs to build branded cover + CV PDFs
    _render_cv_profiles_section(client_name, candidates)

    # Start over
    st.divider()
    if st.button("Create another shortlist"):
        for key in list(st.session_state.keys()):
            if key.startswith("sl_") or key.startswith("cand_") or key.startswith("career_") or key.startswith("proofread"):
                del st.session_state[key]
        st.rerun()


def _render_cv_profiles_section(client_name: str, candidates: list[dict]):
    """Reuse the CVs already uploaded for the shortlist to build branded CV profile PDFs."""
    import zipfile

    form_section("CV Profiles (branded cover + CV)")

    eligible = [c for c in candidates if c.get("cv_bytes") and c.get("name", "").strip()]
    if not eligible:
        st.caption("No CVs available — manual entries can't be turned into CV profiles.")
        return

    st.caption(
        f"Reuses the {len(eligible)} CV(s) already uploaded above. "
        f"Each candidate gets a branded cover page (\"CV of {{name}} prepared for {client_name}\") "
        f"merged with their CV into one PDF."
    )

    if not client_name.strip():
        st.warning("Client name is empty — set it in step 1 to enable CV Profiles.")
        return

    if st.button("Create CV Profiles", type="primary", key="sl_cvp_build"):
        from views.cv_profiles import build_profiles_from_items

        items = [
            {
                "name": c["name"].strip(),
                "cv_bytes": c["cv_bytes"],
                "cv_filename": c.get("source_file", ""),
            }
            for c in eligible
        ]
        st.session_state.sl_cvp_results = build_profiles_from_items(items, client_name.strip())
        st.rerun()

    results = st.session_state.get("sl_cvp_results")
    if not results:
        return

    files = results.get("files", {})
    errors = results.get("errors", [])

    for err in errors:
        st.error(err)

    if not files:
        st.warning("No CV profiles were produced.")
        return

    st.success(f"Built {len(files)} CV profile{'s' if len(files) != 1 else ''}.")

    for fname, fbytes in files.items():
        st.download_button(
            label=f"Download {fname}",
            data=fbytes,
            file_name=fname,
            mime="application/pdf",
            key=f"sl_cvp_dl_{fname}",
        )

    if len(files) > 1:
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for fname, fbytes in files.items():
                zf.writestr(fname, fbytes)
        st.download_button(
            label="Download all CV profiles (.zip)",
            data=zip_buf.getvalue(),
            file_name=f"CV Profiles - {client_name}.zip",
            mime="application/zip",
            type="primary",
            key="sl_cvp_dl_zip",
        )
