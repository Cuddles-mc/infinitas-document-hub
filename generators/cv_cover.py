"""CV Cover Page generator — fills the Infinitas CV Cover Template.

The template ships with two placeholder paragraphs:
    "CV OF "         → append the candidate name
    "Prepared for "  → append the client company name

The appended text inherits the run formatting of the existing placeholder
run, so the big blue cover-page heading style is preserved.
"""

import copy
import io
import os

from docx import Document


TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "assets", "cv_cover", "CV Cover Template.docx"
)


def generate_cover_docx(candidate_name: str, company_name: str) -> bytes:
    """Return .docx bytes for a filled CV cover page."""
    doc = Document(TEMPLATE_PATH)

    filled_name = _append_after_placeholder(doc, "CV OF", candidate_name.upper())
    filled_company = _append_after_placeholder(doc, "Prepared for", company_name)

    if not filled_name:
        raise ValueError("Could not find 'CV OF' placeholder in template.")
    if not filled_company:
        raise ValueError("Could not find 'Prepared for' placeholder in template.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _append_after_placeholder(doc, placeholder: str, text: str) -> bool:
    """Append `text` to the first paragraph starting with `placeholder`.

    Copies the last run's formatting (rPr) so the new text matches the
    heading style of the placeholder.
    """
    for p in doc.paragraphs:
        if p.text.strip().startswith(placeholder):
            last_run = p.runs[-1] if p.runs else None
            new_run = p.add_run(text)
            if last_run is not None and last_run._element.rPr is not None:
                donor_rPr = copy.deepcopy(last_run._element.rPr)
                if new_run._element.rPr is not None:
                    new_run._element.remove(new_run._element.rPr)
                new_run._element.insert(0, donor_rPr)
            return True
    return False
