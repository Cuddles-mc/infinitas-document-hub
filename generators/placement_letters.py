"""Placement letter document generator.

Takes a data dict with placement details, returns branded .docx bytes
for Client Confirmation and/or Candidate Confirmation letters.

Header: Infinitas logo right-aligned in Word header.
Footer: Infinitas logo centred + address text in Word footer.
Matches reference files exactly.
"""

import io
import os
import platform
from datetime import datetime

from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
ASSETS_DIR = os.path.join(os.path.dirname(__file__), "..", "assets", "placement")

CONSULTANT_DETAILS = {
    "Jason Beith": {"key": "jason", "title": "Director"},
    "Tate McClenaghan": {"key": "tate", "title": "Partner"},
    "Kelsi Flynn": {"key": "kelsi", "title": "Senior Consultant"},
}

# Colours
PRIMARY_BLUE = RGBColor(0x00, 0x48, 0x99)
DARK_NAVY = RGBColor(0x0E, 0x28, 0x41)
BODY_TEXT = RGBColor(0x37, 0x41, 0x51)
FOOTER_GREY = RGBColor(0x76, 0x76, 0x76)

# Font settings — Aptos on Windows, Calibri on Linux (Carlito is metric-compatible)
FONT_NAME = "Aptos" if platform.system() == "Windows" else "Calibri"
FONT_SIZE = Pt(10.5)
FOOTER_FONT_SIZE = Pt(8)

# Page margins (matching reference files)
MARGIN_TOP = Emu(698500)     # ~1.94cm
MARGIN_SIDES = Emu(762000)   # ~2.12cm

# Image sizes from reference
HEADER_LOGO_WIDTH = Emu(1714500)   # ~4.76cm
FOOTER_LOGO_WIDTH = Emu(1428750)   # ~3.97cm

# Header/footer distance from page edge
HEADER_FOOTER_DISTANCE = Emu(449580)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _get_logo_path():
    """Get path to the Infinitas logo."""
    return os.path.join(ASSETS_DIR, "header_logo.png")


def _set_run_font(run, font_name=FONT_NAME, font_size=FONT_SIZE, color=BODY_TEXT, bold=False):
    """Apply consistent font formatting to a run."""
    run.font.name = font_name
    run.font.size = font_size
    run.font.color.rgb = color
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_paragraph(doc, text="", color=BODY_TEXT, bold=False, space_after=None, space_before=None):
    """Add a paragraph with consistent formatting."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    if space_after is not None:
        pf.space_after = space_after
    if space_before is not None:
        pf.space_before = space_before
    if text:
        run = para.add_run(text)
        _set_run_font(run, color=color, bold=bold)
    return para


def _set_document_defaults(doc):
    """Set default font for the entire document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    font.color.rgb = BODY_TEXT
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _set_page_layout(doc):
    """Set page margins, header/footer distance, vertical centering."""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_SIDES
        section.left_margin = MARGIN_SIDES
        section.right_margin = MARGIN_SIDES
        section.header_distance = HEADER_FOOTER_DISTANCE
        section.footer_distance = HEADER_FOOTER_DISTANCE
        # Vertically centre body content on page
        sectPr = section._sectPr
        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
        sectPr.append(vAlign)


def _add_blue_border(para, edge="bottom"):
    """Add a blue (#004899) single border to a paragraph edge."""
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:{edge} w:val="single" w:sz="8" w:space="1" w:color="004899"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_header(doc):
    """Add branded header: logo right-aligned, blue rule underneath."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Logo paragraph — right-aligned
    logo_para = header.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = logo_para.add_run()
    run.add_picture(_get_logo_path(), width=HEADER_LOGO_WIDTH)

    # Empty paragraph with blue bottom border (the blue line under header)
    rule_para = header.add_paragraph()
    _add_blue_border(rule_para, "bottom")


def _add_footer(doc):
    """Add branded footer: blue rule on top, centred logo, address, URL."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Empty paragraph with blue top border (the blue line above footer)
    rule_para = footer.paragraphs[0]
    rule_para.text = ""
    _add_blue_border(rule_para, "top")

    # Logo — centred
    logo_para = footer.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run()
    run.add_picture(_get_logo_path(), width=FOOTER_LOGO_WIDTH)

    # Address line — grey, 8pt, centred
    addr_para = footer.add_paragraph()
    addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = addr_para.add_run("Infinitas Talent Limited, PO BOX 357, Shortland Street, Auckland, 1140")
    _set_run_font(run, font_size=FOOTER_FONT_SIZE, color=FOOTER_GREY)

    # URL — grey, 8pt, centred
    url_para = footer.add_paragraph()
    url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = url_para.add_run("https://infinitas.co.nz")
    _set_run_font(run, font_size=FOOTER_FONT_SIZE, color=FOOTER_GREY)


def _add_address_block(doc, lines):
    """Add address as separate paragraphs per line (matching reference style)."""
    for line in lines:
        _add_paragraph(doc, line, space_after=Pt(0), space_before=Pt(0))


def _add_details_table(doc, rows_data):
    """Add a bordered table with blue labels (matching reference style)."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set table-level borders: thin single lines all around and between cells
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    # Remove any existing borders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    for row_idx, (label, value) in enumerate(rows_data):
        row = table.rows[row_idx]

        # Label cell — blue text
        label_cell = row.cells[0]
        label_cell.width = Emu(2122 * 635)
        label_cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        label_cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        label_run = label_cell.paragraphs[0].add_run(f" {label}")
        _set_run_font(label_run, color=PRIMARY_BLUE, bold=False, font_size=FONT_SIZE)

        # Value cell — body text colour
        value_cell = row.cells[1]
        value_cell.width = Emu(7097 * 635)
        value_cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        value_cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        value_run = value_cell.paragraphs[0].add_run(value)
        _set_run_font(value_run, color=BODY_TEXT, bold=False, font_size=FONT_SIZE)

    return table


def _get_signature_path(consultant_key):
    """Get path to consultant's signature image, or None if not available."""
    path = os.path.join(ASSETS_DIR, f"{consultant_key}_signature.png")
    return path if os.path.exists(path) else None


def _add_sign_off(doc, consultant_name, consultant_title):
    """Add sign-off block: signature image (or blank lines), name in bold navy, title in blue."""
    _add_paragraph(doc, "Yours sincerely,", space_after=Pt(0), space_before=Pt(12))

    # Signature image or blank lines
    details = CONSULTANT_DETAILS.get(consultant_name, {"key": "tate", "title": "Partner"})
    sig_path = _get_signature_path(details["key"])
    if sig_path:
        sig_para = doc.add_paragraph()
        sig_para.paragraph_format.space_before = Pt(4)
        sig_para.paragraph_format.space_after = Pt(4)
        run = sig_para.add_run()
        run.add_picture(sig_path, height=Emu(600000))  # ~1.67cm tall
    else:
        for _ in range(3):
            _add_paragraph(doc, "", space_after=Pt(0), space_before=Pt(0))

    # Name — bold, dark navy
    _add_paragraph(doc, consultant_name, color=DARK_NAVY, bold=True,
                   space_after=Pt(0), space_before=Pt(0))
    # Title — blue
    _add_paragraph(doc, consultant_title, color=PRIMARY_BLUE, bold=False,
                   space_after=Pt(0), space_before=Pt(0))
    # Company — body text
    _add_paragraph(doc, "Infinitas Talent", space_after=Pt(0), space_before=Pt(0))


def _format_date(d):
    """Format a date object as '4 March 2026' (no leading zero)."""
    if isinstance(d, str):
        return d
    return f"{d.day} {d.strftime('%B')} {d.year}"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def generate_client_letter(data: dict) -> bytes:
    """Generate the Client Confirmation letter and return .docx bytes."""
    consultant_name = data["consultant"]
    details = CONSULTANT_DETAILS.get(consultant_name, {"key": "tate", "title": "Partner"})
    consultant_title = details["title"]

    candidate = data["candidate_name"]
    company = data["client_company"]
    client_name = data["client_contact_name"]
    client_title = data.get("client_contact_title", "")
    position = data["position"]
    letter_date = data.get("letter_date") or _format_date(datetime.now())

    doc = Document()
    _set_document_defaults(doc)
    _set_page_layout(doc)
    _add_header(doc)
    _add_footer(doc)

    # Date
    _add_paragraph(doc, letter_date, space_after=Pt(12), space_before=Pt(0))

    # Client address block — separate paragraphs per line
    address_lines = []
    if client_title:
        address_lines.append(f"{client_name} \u2013 {client_title}")
    else:
        address_lines.append(client_name)
    address_lines.append(company)
    client_address = data.get("client_address", "")
    if client_address:
        for line in client_address.split("\n"):
            stripped = line.strip()
            if stripped:
                address_lines.append(stripped)
    _add_address_block(doc, address_lines)

    # Salutation
    client_first = client_name.split()[0] if client_name else client_name
    _add_paragraph(doc, f"Dear {client_first},", space_after=Pt(12), space_before=Pt(12))

    # Body
    _add_paragraph(doc, "Thank you for partnering with Infinitas Talent on this search.",
                   space_after=Pt(12))

    _add_paragraph(
        doc,
        f"We are delighted to confirm that {candidate} has accepted the "
        f"position of {position} at {company}.",
        space_after=Pt(12),
    )

    _add_paragraph(doc, "The details of the placement are outlined below.",
                   space_after=Pt(12))

    # Details table
    table_rows = [
        ("Start Date", data.get("start_date", "")),
        ("Position Title", position),
        ("Salary", data.get("salary", "")),
        ("Hiring Manager", data.get("hiring_manager", "")),
        ("Location of Work", data.get("location_of_work", "As agreed")),
        ("Guarantee Period", data.get("guarantee_period", "3 months")),
    ]
    _add_details_table(doc, table_rows)

    # Closing
    _add_paragraph(
        doc,
        f"{client_first}, it has been a real pleasure working with you throughout this "
        f"process. I am confident that {candidate} will be an excellent addition to the "
        f"team at {company}, and I look forward to hearing how they settle in.",
        space_after=Pt(12),
        space_before=Pt(12),
    )

    _add_paragraph(
        doc,
        "We are always here to help, so please do reach out if there is anything "
        "further we can assist with.",
        space_after=Pt(0),
    )

    # Sign-off
    _add_sign_off(doc, consultant_name, consultant_title)

    # Return bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def generate_candidate_letter(data: dict) -> bytes:
    """Generate the Candidate Confirmation letter and return .docx bytes."""
    consultant_name = data["consultant"]
    details = CONSULTANT_DETAILS.get(consultant_name, {"key": "tate", "title": "Partner"})
    consultant_title = details["title"]

    candidate = data["candidate_name"]
    company = data["client_company"]
    position = data["position"]
    letter_date = data.get("letter_date") or _format_date(datetime.now())

    doc = Document()
    _set_document_defaults(doc)
    _set_page_layout(doc)
    _add_header(doc)
    _add_footer(doc)

    # Date
    _add_paragraph(doc, letter_date, space_after=Pt(12), space_before=Pt(0))

    # Candidate address block — separate paragraphs per line
    address_lines = [candidate]
    candidate_address = data.get("candidate_address", "")
    if candidate_address:
        for line in candidate_address.split("\n"):
            stripped = line.strip()
            if stripped:
                address_lines.append(stripped)
    _add_address_block(doc, address_lines)

    # Salutation
    candidate_first = candidate.split()[0] if candidate else candidate
    _add_paragraph(doc, f"Dear {candidate_first},", space_after=Pt(12), space_before=Pt(12))

    # Body
    _add_paragraph(doc, "Congratulations on your new role.", space_after=Pt(12))

    _add_paragraph(
        doc,
        f"We are delighted to confirm that you have accepted the position of "
        f"{position} at {company}.",
        space_after=Pt(12),
    )

    _add_paragraph(doc, "The details of your new role are outlined below.",
                   space_after=Pt(12))

    # Details table (no Guarantee Period for candidate)
    table_rows = [
        ("Start Date", data.get("start_date", "")),
        ("Position Title", position),
        ("Salary", data.get("salary", "")),
        ("Hiring Manager", data.get("hiring_manager", "")),
        ("Location of Work", data.get("location_of_work", "As agreed")),
    ]
    _add_details_table(doc, table_rows)

    # Closing
    _add_paragraph(
        doc,
        f"{candidate_first}, it has been a real pleasure getting to know you throughout "
        f"this process. I look forward to staying in touch as you settle into the role. "
        f"We are always here to support you, so please reach out any time if there is "
        f"anything we can help with.",
        space_after=Pt(12),
        space_before=Pt(12),
    )

    _add_paragraph(doc, "All the very best and congratulations again.",
                   space_after=Pt(0))

    # Sign-off
    _add_sign_off(doc, consultant_name, consultant_title)

    # Return bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
