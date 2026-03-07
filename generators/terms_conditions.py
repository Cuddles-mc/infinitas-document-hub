"""Terms & Conditions document generator.

Takes a data dict with service type toggles and fee details,
returns branded .docx bytes with only the relevant clauses.
"""

import io
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "terms-conditions.docx"

# Clause heading text -> clause number in the original template
CLAUSE_HEADINGS = {
    "DEFINITIONS AND INTERPRETATION": 1,
    "TERM": 2,
    "RIGHT OF RENEWAL": 3,
    "PLACEMENT FEE": 4,
    "LIABILITY TO PAY A PLACEMENT FEE": 5,
    "CONTRACTOR OR TEMPORARY WORKER SERVICES": 6,
    "FEES FOR FURTHER CONTRACTING": 7,
    "RETAINED ASSIGNMENT AND EXECUTIVE SEARCH": 8,
    "EXPENSES": 9,
    "PLACEMENT GUARANTEE": 10,
    "CLIENT OBLIGATIONS": 11,
    "INFINITAS TALENT OBLIGATIONS": 12,
    "LIMITATION OF LIABILITY": 13,
    "GST": 14,
    "ANTI-CORRUPTION": 15,
    "CONFIDENTIALITY AND PRIVACY": 16,
    "TERMINATION": 17,
    "CONSEQUENCES OF EXPIRY OR TERMINATION": 18,
    "GENERAL PROVISIONS": 19,
    "SCHEDULE 1": 20,
}

# Which clauses to remove per service type toggle
REMOVABLE = {
    "perm": [4, 5],       # Placement Fee + Liability to Pay
    "contract": [6, 7],   # Contractor/Temp Fees + Further Contracting
    "exec": [8],          # Retained/Exec Search Fees
}

# Guarantee period text variants
GUARANTEE_TEXT = {
    3: "three (3) calendar months",
    6: "six (6) calendar months",
    12: "twelve (12) calendar months",
}

FONT_NAME = "Aptos"
FONT_SIZE = Pt(10)


def _find_heading_ranges(doc):
    """Map each Heading 1 paragraph to its clause number and index range."""
    ranges = []
    heading_indices = []

    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Heading 1":
            text = para.text.strip().upper()
            for key, num in CLAUSE_HEADINGS.items():
                if key in text:
                    heading_indices.append((num, i))
                    break

    for idx, (num, start) in enumerate(heading_indices):
        if idx + 1 < len(heading_indices):
            end = heading_indices[idx + 1][1]
        else:
            end = len(doc.paragraphs)
        ranges.append((num, start, end))

    return ranges


def _remove_paragraphs(doc, indices_to_remove: set):
    """Remove paragraphs by index from the document XML."""
    body = doc.element.body
    to_remove = []
    for i, para in enumerate(doc.paragraphs):
        if i in indices_to_remove:
            to_remove.append(para._element)
    for elem in to_remove:
        body.remove(elem)


def _build_clause_map(removed_clauses: set) -> dict:
    """Build mapping of old clause number -> new clause number."""
    all_clauses = list(range(1, 21))
    remaining = [c for c in all_clauses if c not in removed_clauses]
    return {old: new for new, old in enumerate(remaining, 1)}


def _update_cross_references(doc, clause_map: dict):
    """Update all 'clause N' references in the document."""
    pattern = re.compile(r'(clause[s]?\s+)(\d+)(\.\d+)?', re.IGNORECASE)

    def replace_ref(match):
        prefix = match.group(1)
        old_num = int(match.group(2))
        sub = match.group(3) or ""
        new_num = clause_map.get(old_num, old_num)
        return f"{prefix}{new_num}{sub}"

    for para in doc.paragraphs:
        if "clause" in para.text.lower():
            for run in para.runs:
                if "clause" in run.text.lower():
                    run.text = pattern.sub(replace_ref, run.text)


def _update_guarantee_definition(doc, months: int):
    """Update the Guarantee Period definition text."""
    old = "three (3) calendar months"
    new = GUARANTEE_TEXT.get(months, f"{months} calendar months")
    for para in doc.paragraphs:
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    return
            # Fallback: spans runs
            full = para.text.replace(old, new)
            for run in para.runs[1:]:
                run.text = ""
            para.runs[0].text = full
            return


def _fill_client_name(doc, client_name: str):
    """Replace [PARTY 2] in Table 1 with the client company name."""
    if len(doc.tables) >= 2:
        cell = doc.tables[1].rows[0].cells[1]
        for para in cell.paragraphs:
            for run in para.runs:
                if "[PARTY 2]" in run.text:
                    run.text = run.text.replace("[PARTY 2]", client_name)
                    return
            if "[PARTY 2]" in para.text:
                para.text = para.text.replace("[PARTY 2]", client_name)
                return


def _set_run_font(run):
    """Apply consistent font formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _rewrite_schedule_1(doc, data: dict):
    """Rewrite the Schedule 1 fee text based on enabled service types.

    Removes all content after the Schedule 1 heading from the XML,
    then appends fresh, properly-formatted paragraphs.
    """
    # Find Schedule 1 heading element
    sched_element = None
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1" and "SCHEDULE" in para.text.upper():
            sched_element = para._element
            break

    if sched_element is None:
        return

    # Remove all sibling elements after the schedule heading
    body = doc.element.body
    to_remove = []
    found = False
    for child in body:
        if child is sched_element:
            found = True
            continue
        if found:
            to_remove.append(child)
    for elem in to_remove:
        body.remove(elem)

    # Build fee entry text blocks
    entries = _build_schedule_entries(data)

    # Append fresh paragraphs with proper formatting
    for text in entries:
        para = doc.add_paragraph()
        run = para.add_run(text)
        _set_run_font(run)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.space_before = Pt(4)


def _build_schedule_entries(data: dict) -> list[str]:
    """Build the list of fee text entries for Schedule 1."""
    entries = []

    if data.get("perm_enabled", True):
        pct = data.get("perm_fee_pct", 18)
        basis = data.get("perm_basis", "total salary package")
        structure = data.get("perm_structure", "retained")
        guarantee = data.get("guarantee_months", 3)

        if structure == "fixed_fee":
            fee_text = f"Permanent Fees: A fixed fee of {data.get('perm_fixed_fee', 'TBC')}."
        else:
            fee_text = f"Permanent Fees: {pct}% based on the candidate\u2019s {basis}."

        if structure == "retained":
            fee_text += (
                " Unless otherwise stated, Infinitas Talent uses an industry standard "
                "retained fee structure and is invoiced in three instalments. One third on "
                "acceptance of the assignment, one third on presentation of the shortlist "
                f"and one third on placement. The placement guarantee is {guarantee} months."
            )
        elif structure == "contingent":
            fee_text += (
                f" Invoiced on placement. The placement guarantee is {guarantee} months."
            )

        entries.append(fee_text)

        # Fixed term
        ft_text = (
            f"Fixed Term Contract Fees: {pct}% Fees will be calculated on the candidates "
            f"{basis} Pro-Rata for the length of the fixed term placement (calculated in months). "
            "The minimum fee period to engage a fixed term candidate is six months. "
            f"The fixed term placement guarantee is {guarantee} months."
        )
        entries.append(ft_text)

    if data.get("contract_enabled", False):
        margin = data.get("contract_margin_pct", 25)
        entries.append(
            f"Contracting Fees: Infinitas Talent charges contract fees on either an hourly "
            f"or daily basis as agreed with you the client. Margin percentages on contracting "
            f"assignments are {margin}%."
        )

    if data.get("exec_enabled", False):
        pct = data.get("exec_fee_pct", 25)
        basis = data.get("exec_basis", "total salary package")

        if data.get("exec_structure") == "fixed_fee":
            exec_text = f"Executive Search Fees: A fixed fee of {data.get('exec_fixed_fee', 'TBC')}."
        else:
            exec_text = (
                f"Executive Search Fees: For Executive Recruitment Infinitas Talent\u2019s "
                f"fee is {pct}% of the candidate\u2019s {basis}. Infinitas Talent uses an "
                "industry standard retained fee structure and is invoiced in three instalments. "
                "One third on acceptance of the assignment, one third on presentation of the "
                "shortlist and one third on placement."
            )
        exec_text += (
            "\n\nExecutive Search Recruitment is defined as senior/executive leadership or "
            "specialised positions where a customised advertising and/or search process is "
            "undertaken on an exclusive basis."
        )
        entries.append(exec_text)

    # Contract buy-out
    if data.get("contract_enabled", False) or data.get("perm_enabled", True):
        entries.append(
            "Contract Buy out: For temporary or contract candidates, who are offered "
            "permanent positions with the client, a Pro-Rata fee will be calculated for a "
            "period of up to twelve months, with a minimum period of six months to be charged "
            "on acceptance of the engagement or employment by the client."
        )

    # GST note
    entries.append(
        "All Fees quoted exclude \u201cGST\u201d Goods and Services Tax. GST will be "
        "added to final invoices sent out by Infinitas Talent."
    )

    return entries


def _add_signature_block(doc, include_infinitas: bool, include_client: bool, adobe_sign: bool):
    """Append a styled signature block table at the end of the document."""
    if not include_infinitas and not include_client:
        return

    doc.add_paragraph("")  # spacer

    rows_needed = 0
    if include_infinitas:
        rows_needed += 1
    if include_client:
        rows_needed += 1

    table = doc.add_table(rows=rows_needed * 2, cols=3)

    # Style the table with blue top/bottom borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="004899"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="004899"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    row_idx = 0
    if include_infinitas:
        cells = table.rows[row_idx].cells
        cells[0].text = "Signature" if not adobe_sign else "{{Sig_es_:signer1:signature}}"
        cells[1].text = "Name"
        cells[2].text = "Date" if not adobe_sign else "{{Dte_es_:signer1:date}}"
        label_cells = table.rows[row_idx + 1].cells
        label_cells[0].text = "Signed for and on behalf of Infinitas Talent Limited"
        row_idx += 2

    if include_client:
        signer = "signer2" if include_infinitas else "signer1"
        cells = table.rows[row_idx].cells
        cells[0].text = "Signature" if not adobe_sign else f"{{{{Sig_es_:{signer}:signature}}}}"
        cells[1].text = "Name"
        cells[2].text = "Date" if not adobe_sign else f"{{{{Dte_es_:{signer}:date}}}}"
        label_cells = table.rows[row_idx + 1].cells
        label_cells[0].text = "Signed for and on behalf of the Client"

    # Apply font formatting to all cells
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                for run in para.runs:
                    _set_run_font(run)


def generate_docx(data: dict) -> bytes:
    """Generate T&Cs .docx from form data."""
    doc = Document(str(TEMPLATE_PATH))

    # 1. Fill client name
    _fill_client_name(doc, data.get("client_name", ""))

    # 2. Update guarantee definition
    _update_guarantee_definition(doc, data.get("guarantee_months", 3))

    # 3. Determine which clauses to remove
    clauses_to_remove = set()
    if not data.get("perm_enabled", True):
        clauses_to_remove.update(REMOVABLE["perm"])
    if not data.get("contract_enabled", False):
        clauses_to_remove.update(REMOVABLE["contract"])
    if not data.get("exec_enabled", False):
        clauses_to_remove.update(REMOVABLE["exec"])

    # 4. Remove clause paragraphs
    if clauses_to_remove:
        ranges = _find_heading_ranges(doc)
        indices_to_remove = set()
        for clause_num, start, end in ranges:
            if clause_num in clauses_to_remove:
                for i in range(start, end):
                    indices_to_remove.add(i)
        _remove_paragraphs(doc, indices_to_remove)

    # 5. Re-number clauses and update cross-references
    clause_map = _build_clause_map(clauses_to_remove)
    _update_cross_references(doc, clause_map)

    # 6. Rewrite Schedule 1
    _rewrite_schedule_1(doc, data)

    # 7. Add signature block
    _add_signature_block(
        doc,
        data.get("sig_infinitas", False),
        data.get("sig_client", False),
        data.get("adobe_sign", False),
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
