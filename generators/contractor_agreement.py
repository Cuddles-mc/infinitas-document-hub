"""Contractor Agreement document generator.

Takes a data dict with Schedule 1 fields, returns branded .docx bytes.
Supports Sole Trader and Limited Company templates.

Schedule 1 is rewritten from scratch as a properly formatted table
rather than using the template's messy inline text layout.
"""

import io
import platform
from pathlib import Path
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


TEMPLATES = {
    "sole_trader": Path(__file__).parent.parent / "templates" / "contractor-agreement-sole-trader.docx",
    "ltd_company": Path(__file__).parent.parent / "templates" / "contractor-agreement-ltd-company.docx",
}

# Brand constants
FONT_NAME = "Aptos" if platform.system() == "Windows" else "Calibri"
FONT_SIZE = Pt(10)
PRIMARY_BLUE = RGBColor(0x00, 0x48, 0x99)
BODY_TEXT = RGBColor(0x37, 0x41, 0x51)
LABEL_COLOUR = PRIMARY_BLUE

# DocuSign text tag patterns (invisible white text picked up by DocuSign)
DOCUSIGN_TAGS = {
    "sig_1": "\\s1\\",
    "name_1": "\\n1\\",
    "date_1": "\\d1\\",
    "sig_2": "\\s2\\",
    "name_2": "\\n2\\",
    "date_2": "\\d2\\",
}


def _set_run_font(run, font_name=FONT_NAME, font_size=FONT_SIZE, color=BODY_TEXT, bold=False):
    """Apply consistent font formatting to a run."""
    run.font.name = font_name
    run.font.size = font_size
    run.font.color.rgb = color
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_paragraph(doc, text="", color=BODY_TEXT, bold=False, space_after=Pt(2), space_before=Pt(2)):
    """Add a paragraph with consistent formatting."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    if text:
        run = para.add_run(text)
        _set_run_font(run, color=color, bold=bold)
    return para


def _add_schedule_table(doc, rows_data: list[tuple[str, str]]):
    """Add a formatted Schedule 1 details table.

    rows_data: list of (label, value) tuples
    """
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="004899"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="004899"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    for row_idx, (label, value) in enumerate(rows_data):
        row = table.rows[row_idx]

        # Label cell — blue text
        label_cell = row.cells[0]
        label_cell.width = Emu(2500 * 635)
        label_cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        label_cell.paragraphs[0].paragraph_format.space_after = Pt(3)
        label_run = label_cell.paragraphs[0].add_run(label)
        _set_run_font(label_run, color=LABEL_COLOUR, bold=False)

        # Value cell
        value_cell = row.cells[1]
        value_cell.width = Emu(6500 * 635)
        value_cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        value_cell.paragraphs[0].paragraph_format.space_after = Pt(3)
        value_run = value_cell.paragraphs[0].add_run(value or "")
        _set_run_font(value_run, color=BODY_TEXT, bold=False)

    return table


def _add_signature_block(doc, label: str, signer_num: int, docusign: bool = False):
    """Add a signature block with optional DocuSign text tags."""
    _add_paragraph(doc, "", space_after=Pt(6))

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Style
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="004899"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # Header row: Signature | Name | Date
    headers = ["Signature", "Name", "Date"]
    for ci, header_text in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        run = cell.paragraphs[0].add_run(header_text)
        _set_run_font(run, color=LABEL_COLOUR, bold=True)

        if docusign:
            # Add invisible DocuSign text tag
            tag_key = ["sig", "name", "date"][ci]
            tag = DOCUSIGN_TAGS.get(f"{tag_key}_{signer_num}", "")
            if tag:
                tag_run = cell.paragraphs[0].add_run(f" {tag}")
                _set_run_font(tag_run, color=RGBColor(0xFF, 0xFF, 0xFF), font_size=Pt(1))

    # Spacer row for actual signature
    for ci in range(3):
        cell = table.rows[1].cells[ci]
        cell.paragraphs[0].paragraph_format.space_before = Pt(20)
        cell.paragraphs[0].paragraph_format.space_after = Pt(4)

    # Label below table
    _add_paragraph(doc, label, color=BODY_TEXT, bold=False, space_after=Pt(8))


def _find_schedule_start(doc) -> int | None:
    """Find the paragraph index where Schedule 1 content starts.

    Looks for 'Date of Agreement:' which is the first field in Schedule 1.
    """
    for i, para in enumerate(doc.paragraphs):
        if "Date of Agreement:" in para.text:
            return i
    return None


def _remove_from_index(doc, start_idx: int):
    """Remove all paragraph and table elements from start_idx to end of document body."""
    body = doc.element.body

    # Collect elements to remove (paragraphs from start_idx onwards + any tables/elements after)
    start_element = doc.paragraphs[start_idx]._element

    to_remove = []
    found = False
    for child in body:
        if child is start_element:
            found = True
        if found:
            to_remove.append(child)

    for elem in to_remove:
        body.remove(elem)


def _write_sole_trader_schedule(doc, data: dict):
    """Write a clean Schedule 1 for Sole Trader agreements."""
    _add_paragraph(doc, "", space_after=Pt(4))
    _add_paragraph(doc, "SCHEDULE 1", color=PRIMARY_BLUE, bold=True, space_after=Pt(8))

    # Agency details (static)
    agency_rows = [
        ("Date of Agreement", data.get("date_of_agreement", "")),
        ("Recruitment Agency", "Infinitas Talent Limited"),
        ("Address", "2 Princes Street, Auckland Central, Auckland 1010"),
        ("GST/IRD Number", "125-826-008"),
        ("Company No.", "6829273"),
    ]
    _add_schedule_table(doc, agency_rows)

    _add_paragraph(doc, "", space_after=Pt(8))

    # Assignment details
    _add_paragraph(doc, "Assignment Details", color=PRIMARY_BLUE, bold=True, space_after=Pt(4))
    assignment_rows = [
        ("Nominated Client", data.get("nominated_client", "")),
        ("Role", data.get("role", "")),
        ("Commencement Date", data.get("commencement_date", "")),
        ("End Date", data.get("end_date", "")),
        ("Hours of Work", data.get("hours_of_work", "")),
        ("Contract Rate", data.get("contract_rate", "")),
        ("Notice Period", data.get("notice_period", "")),
        ("Other/Travel Expenses", data.get("travel_expenses", "Upon authorisation by the Nominated Client")),
    ]
    _add_schedule_table(doc, assignment_rows)

    # Signature blocks
    docusign = data.get("docusign", False)
    _add_signature_block(doc, "Signed for and on behalf of the Recruitment Agency", 1, docusign)
    _add_signature_block(doc, "Signed for and on behalf of the Contractor", 2, docusign)

    # Footer
    _add_paragraph(doc, "", space_after=Pt(12))
    para = _add_paragraph(
        doc,
        "Infinitas Talent Limited, PO BOX 357, Shortland Street, Auckland, 1140",
        color=RGBColor(0x76, 0x76, 0x76),
        space_after=Pt(0),
    )
    para.runs[0].font.size = Pt(8)
    para2 = _add_paragraph(
        doc,
        "https://infinitas.co.nz",
        color=RGBColor(0x76, 0x76, 0x76),
        space_after=Pt(0),
    )
    para2.runs[0].font.size = Pt(8)


def _write_ltd_company_schedule(doc, data: dict):
    """Write a clean Schedule 1 for Ltd Company agreements."""
    _add_paragraph(doc, "", space_after=Pt(4))
    _add_paragraph(doc, "SCHEDULE 1", color=PRIMARY_BLUE, bold=True, space_after=Pt(8))

    # Principal details (static)
    principal_rows = [
        ("Date of Agreement", data.get("date_of_agreement", "")),
        ("Principal", "Infinitas Talent Limited"),
        ("Address", "2 Princes Street, Auckland Central, Auckland 1010"),
        ("GST/IRD Number", "125-826-008"),
        ("Company No./NZBN", "6829273"),
    ]
    _add_schedule_table(doc, principal_rows)

    _add_paragraph(doc, "", space_after=Pt(8))

    # Provider (contractor) details
    _add_paragraph(doc, "Provider (Contractor)", color=PRIMARY_BLUE, bold=True, space_after=Pt(4))
    provider_rows = [
        ("Provider Company", data.get("provider_company", "")),
        ("Trading As", data.get("trading_as", "")),
        ("Registered Address", data.get("registered_address", "")),
        ("Company No./NZBN", data.get("company_nzbn", "")),
        ("Individual Contractor", data.get("individual_contractor", "")),
        ("IRD Number", data.get("ird_number", "")),
        ("GST Registered", "Yes" if data.get("gst_registered") else "No"),
        ("GST Number", data.get("gst_number", "")),
        ("Bank Account", data.get("bank_account", "")),
    ]
    _add_schedule_table(doc, provider_rows)

    _add_paragraph(doc, "", space_after=Pt(8))

    # Assignment details
    _add_paragraph(doc, "Assignment Details", color=PRIMARY_BLUE, bold=True, space_after=Pt(4))
    assignment_rows = [
        ("Nominated Client", data.get("nominated_client", "")),
        ("Role", data.get("role", "")),
        ("Commencement Date", data.get("commencement_date", "")),
        ("End Date", data.get("end_date", "")),
        ("Hours of Work", data.get("hours_of_work", "")),
        ("Contract Rate", data.get("contract_rate", "")),
        ("Notice Period", data.get("notice_period", "")),
        ("Other/Travel Expenses", data.get("travel_expenses", "Upon authorisation by the Nominated Client")),
    ]
    _add_schedule_table(doc, assignment_rows)

    # Signature blocks
    docusign = data.get("docusign", False)
    _add_signature_block(doc, "Signed for and on behalf of the Principal", 1, docusign)
    _add_signature_block(doc, "Signed for and on behalf of the Provider/Nominated Consultant", 2, docusign)

    # Footer
    _add_paragraph(doc, "", space_after=Pt(12))
    para = _add_paragraph(
        doc,
        "Infinitas Talent Limited, PO BOX 357, Shortland Street, Auckland, 1140",
        color=RGBColor(0x76, 0x76, 0x76),
        space_after=Pt(0),
    )
    para.runs[0].font.size = Pt(8)
    para2 = _add_paragraph(
        doc,
        "https://infinitas.co.nz",
        color=RGBColor(0x76, 0x76, 0x76),
        space_after=Pt(0),
    )
    para2.runs[0].font.size = Pt(8)


def generate_docx(data: dict) -> bytes:
    """Generate a Contractor Agreement .docx from form data.

    data must include:
        contractor_type: "sole_trader" or "ltd_company"
        Plus all Schedule 1 fields.
        Optional: docusign: bool — adds invisible DocuSign text tags to signature blocks.
    Returns .docx file as bytes.
    """
    ctype = data.get("contractor_type", "sole_trader")
    template_path = TEMPLATES[ctype]
    doc = Document(str(template_path))

    # Find and remove old Schedule 1 content
    sched_start = _find_schedule_start(doc)
    if sched_start is not None:
        _remove_from_index(doc, sched_start)

    # Write new formatted Schedule 1
    if ctype == "sole_trader":
        _write_sole_trader_schedule(doc, data)
    else:
        _write_ltd_company_schedule(doc, data)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
