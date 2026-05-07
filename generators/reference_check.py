"""Reference check document generator (v2 — Option 8 hand-refined spec).

Builds the document from scratch via python-docx. No .dotx template required.

Spec source: handovers/2026-05-06-handover-reference-template-redesign.md

Design summary:
- Top banner: navy fill, "Reference Check" white left + white logo right.
- Three content tables (Details, Reference Questions, Completed By), each with
  the section header merged into row 0 of the table so the header always
  travels with its content across page breaks.
- Details: 3-column metadata table (label / title / employer). Title and
  employer split into separate cells for the two referee role rows.
- Questions: variable length. Standard 26 questions live in STANDARD_QUESTIONS.
  Callers can pass a custom `questions` list to add or remove questions per
  reference (additional-questions capability — plumbed but not exposed in UI
  yet).
- Sign-off: single-row table with header bar.
- Footer: horizontal logo + address + URL.
"""

import io
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Brand palette ---
NAVY = "0E2841"
BLUE = "004899"
LBL = "6B7280"
BODY = "374151"
SOFT = "F3F4F6"
WHITE = "FFFFFF"
BOX = "B7BFC9"

# --- Asset paths ---
TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
LOGO_FOOTER = TEMPLATES_DIR / "logo-footer.png"
LOGO_WHITE = TEMPLATES_DIR / "logo-white.png"

# --- Standard 26 questions (verbatim Infinitas reference set) ---
STANDARD_QUESTIONS = [
    "What was your relationship to the candidate? Did they report into you?",
    "How long were they with the company? How long did they report to you?",
    "Why did the candidate leave the company?",
    "What were their main role and responsibilities?",
    "In relation to your expectations how would you rate their overall performance?",
    "How would you comment on their technical skills required to carry out their role?",
    "How would you comment on their work ethic?",
    "What are their main strengths?",
    "What are their main areas for development?",
    "What is their ability to meet deadlines?",
    "How would you comment on their decision-making ability?",
    "Did they manage a team? What was their management style / ability?",
    "Do they work well in the team / company with others?",
    "Were there ever any conflicts with any other team members?",
    "How did they manage external stakeholders?",
    "How would you describe their ability to deal with pressure?",
    "How would you describe their ability to cope with change?",
    "Have you ever had to question their integrity or honesty?",
    "How would you comment on their attendance? Any sick leave taken or attendance issues?",
    "Have you ever had to raise any behavioural or performance issues with the candidate? If so, what were the reasons and the outcome?",
    "Are you aware of any health condition or substance dependency or anything else that might affect their ability to do their job?",
    "If you had the opportunity, would you employ them in a similar role? Or at all?",
    "Would you recommend them for the role for which they are applying?",
    "Are we able to share this reference with the candidate?",
    "Are we able to share this reference with the client?",
    "Is there any additional information that you feel we should consider as part of our evaluation as to their suitability for this job or are there any other comments you would like to make?",
]

# Backward-compat alias for callers that imported QUESTIONS
QUESTIONS = STANDARD_QUESTIONS


# --- Low-level XML helpers ---

def _shd(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _borders(cell, top=None, left=None, right=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    edges = {"top": top, "left": left, "bottom": bottom, "right": right}
    for edge, spec in edges.items():
        if spec is None:
            continue
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), spec.get("val", "single"))
        b.set(qn("w:sz"), str(spec.get("sz", 4)))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), spec.get("color", "auto"))
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _cell_margins(cell, top=80, left=120, bottom=80, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:cantSplit")):
        trPr.remove(old)
    trPr.append(OxmlElement("w:cantSplit"))


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _table_full_width(table, total_pct=100):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_pct * 50))
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def _run(p, text, *, font="Arial", size=10, bold=False, color="000000",
         caps=False, italic=False, letter_spacing=None):
    from docx.shared import RGBColor
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rPr.append(rFonts)
    if caps:
        rPr.append(OxmlElement("w:caps"))
    if letter_spacing is not None:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(letter_spacing))
        rPr.append(sp)
    return r


def _set_page_margins(doc, top=0.6, bottom=0.7, left=0.85, right=0.85):
    for s in doc.sections:
        s.top_margin = Inches(top)
        s.bottom_margin = Inches(bottom)
        s.left_margin = Inches(left)
        s.right_margin = Inches(right)


def _split_title_company(s):
    """Heuristic split of 'Title at Company' / 'Title, Company' → (title, company).

    Returns (title, '') if no separator found. Used as a graceful fallback when
    the caller passes the legacy single-field referee_title/referee_previous
    instead of the split fields.
    """
    if not s:
        return ("", "")
    s = s.strip()
    for sep in (" at ", ", "):
        if sep in s:
            t, c = s.split(sep, 1)
            return (t.strip(), c.strip())
    return (s, "")


def _section_header_row(table, row_idx, label, n_cols):
    """Format row 0 of a content table as a blue (#004899) section header bar.

    Merges all cells across n_cols so the header looks like a continuous bar
    spanning the table width. cantSplit + keepNext on the row make sure the
    header always travels with at least one row of content underneath it.
    """
    row = table.rows[row_idx]
    _row_cant_split(row)

    if n_cols > 1:
        first = row.cells[0]
        for i in range(1, n_cols):
            first.merge(row.cells[i])

    cell = row.cells[0]
    _shd(cell, BLUE)
    _cell_margins(cell, top=80, bottom=80, left=140, right=140)
    _borders(
        cell,
        top={"val": "nil"},
        left={"val": "nil"},
        bottom={"val": "nil"},
        right={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    _run(p, label, size=10, bold=True, color=WHITE, caps=True, letter_spacing=8)


# --- Section builders ---

def _build_banner(doc, title="Reference Check"):
    """Top navy banner with white title left + white logo right."""
    hdr = doc.add_table(rows=1, cols=2)
    _table_full_width(hdr)
    _no_table_borders(hdr)
    hdr.columns[0].width = Inches(4.4)
    hdr.columns[1].width = Inches(2.4)

    cell_t = hdr.rows[0].cells[0]
    cell_l = hdr.rows[0].cells[1]
    _shd(cell_t, NAVY)
    _shd(cell_l, NAVY)
    _cell_margins(cell_t, top=180, bottom=180, left=240, right=120)
    _cell_margins(cell_l, top=180, bottom=180, left=120, right=240)
    cell_t.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_l.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    pT = cell_t.paragraphs[0]
    pT.paragraph_format.space_before = Pt(0)
    pT.paragraph_format.space_after = Pt(0)
    _run(pT, title, size=22, bold=True, color=WHITE)

    pL = cell_l.paragraphs[0]
    pL.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pL.paragraph_format.space_before = Pt(0)
    pL.paragraph_format.space_after = Pt(0)
    if LOGO_WHITE.exists():
        pL.add_run().add_picture(str(LOGO_WHITE), width=Inches(1.8))

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(14)


def _build_details(doc, meta):
    """Details table: header row + metadata rows. 3 cols (label / title / company)."""
    title_now, company_now = _split_title_company(meta.get("referee_title", ""))
    if meta.get("referee_current_title"):
        title_now = meta["referee_current_title"]
    if meta.get("referee_current_company"):
        company_now = meta["referee_current_company"]

    title_prev, company_prev = _split_title_company(meta.get("referee_previous", ""))
    if meta.get("referee_previous_title"):
        title_prev = meta["referee_previous_title"]
    if meta.get("referee_previous_company"):
        company_prev = meta["referee_previous_company"]

    rows = [
        ("Reference date",                            meta.get("date", ""),       ""),
        ("Candidate",                                 meta.get("candidate", ""),  ""),
        ("Position applied for",                      meta.get("position", ""),   ""),
        ("Referee",                                   meta.get("referee", ""),    ""),
        ("Current Title and employer",                title_now,                  company_now),
        ("Previous role AND EMployer (If Different)", title_prev,                 company_prev),
    ]
    n_cols = 3
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    _table_full_width(table)
    _no_table_borders(table)
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(2.6)
    table.columns[2].width = Inches(2.0)

    _section_header_row(table, 0, "Details", n_cols)

    for i, (label, val1, val2) in enumerate(rows):
        row = table.rows[i + 1]
        _row_cant_split(row)
        lcell, vcell, ccell = row.cells[0], row.cells[1], row.cells[2]
        _shd(lcell, SOFT)
        for c in (lcell, vcell, ccell):
            _cell_margins(c, top=120, bottom=120, left=160, right=120)
            _borders(
                c,
                top={"sz": 6, "color": BOX},
                left={"sz": 6, "color": BOX},
                bottom={"sz": 6, "color": BOX},
                right={"sz": 6, "color": BOX},
            )
        lp = lcell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        _run(lp, label, size=8.5, color=LBL, caps=True, letter_spacing=40)
        vp = vcell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        _run(vp, val1, size=10.5, color=NAVY, bold=True)
        cp = ccell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(0)
        _run(cp, val2, size=10.5, color=NAVY, bold=True)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(14)


def _build_questions(doc, questions, answers):
    """Q&A table: header row + N boxed question rows. cantSplit per row."""
    n = len(questions)
    table = doc.add_table(rows=1 + n, cols=1)
    _table_full_width(table)
    _no_table_borders(table)

    _section_header_row(table, 0, "Reference Questions", 1)

    for i in range(n):
        row = table.rows[i + 1]
        _row_cant_split(row)
        cell = row.cells[0]
        _cell_margins(cell, top=240, bottom=240, left=200, right=200)
        _borders(
            cell,
            top={"sz": 6, "color": BOX},
            left={"sz": 6, "color": BOX},
            bottom={"sz": 6, "color": BOX},
            right={"sz": 6, "color": BOX},
        )
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)

        qp = cell.add_paragraph()
        qp.paragraph_format.space_before = Pt(0)
        qp.paragraph_format.space_after = Pt(12)
        _run(qp, f"{i + 1}.  ", size=10.5, bold=True, color=BLUE)
        _run(qp, questions[i], size=10.5, bold=True, color=NAVY)

        if isinstance(answers, dict):
            answer = answers.get(str(i), "")
        elif isinstance(answers, (list, tuple)) and i < len(answers):
            answer = answers[i] or ""
        else:
            answer = ""
        if answer.startswith("[GAP] "):
            answer = answer[6:]

        paras = answer.split("\n\n") if answer else [""]
        for j, para_text in enumerate(paras):
            ap = cell.add_paragraph()
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.space_after = Pt(4 if j < len(paras) - 1 else 0)
            ap.paragraph_format.line_spacing = 1.4
            _run(ap, para_text, size=10, color="000000")

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(14)


def _build_signoff(doc, meta):
    """Sign-off table: header row + 1 sign-off row."""
    table = doc.add_table(rows=2, cols=1)
    _table_full_width(table)
    _no_table_borders(table)

    _section_header_row(table, 0, "Completed By", 1)

    row = table.rows[1]
    _row_cant_split(row)
    cell = row.cells[0]
    _cell_margins(cell, top=140, bottom=140, left=160, right=160)
    _borders(
        cell,
        top={"sz": 6, "color": BOX},
        left={"sz": 6, "color": BOX},
        bottom={"sz": 6, "color": BOX},
        right={"sz": 6, "color": BOX},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Reference completed by  ", size=10, color=LBL)
    _run(p, meta.get("completed_by", "Tate McClenaghan"), size=10.5, bold=True, color=NAVY)
    _run(p, "  on  ", size=10, color=LBL)
    _run(p, meta.get("date", ""), size=10.5, bold=True, color=NAVY)


def _build_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    p_logo = footer.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(2)
    if LOGO_FOOTER.exists():
        p_logo.add_run().add_picture(str(LOGO_FOOTER), width=Inches(1.7))

    p_addr = footer.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_addr.paragraph_format.space_before = Pt(0)
    p_addr.paragraph_format.space_after = Pt(0)
    _run(
        p_addr,
        "Infinitas Talent Limited, PO BOX 357, Shortland Street, Auckland, 1140",
        size=9,
        color="6F6F6F",
    )

    p_url = footer.add_paragraph()
    p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_url.paragraph_format.space_before = Pt(0)
    p_url.paragraph_format.space_after = Pt(0)
    _run(p_url, "https://infinitas.co.nz", size=9, color="6F6F6F")


# --- Public API ---

def generate_docx(data: dict, questions=None) -> bytes:
    """Build a branded reference check .docx and return bytes.

    Args:
        data: dict with keys
            candidate_name, position, referee_name,
            referee_title, referee_previous,                  (legacy combined)
            referee_current_title, referee_current_company,   (preferred split)
            referee_previous_title, referee_previous_company, (preferred split)
            reference_date, completed_by,
            answers: dict {str index: str answer} OR list of strings.
        questions: optional list of question strings. Defaults to
            STANDARD_QUESTIONS (the canonical Infinitas 26). Pass a longer
            list to add custom questions per reference. Length must match
            the answer set.

    Returns:
        .docx file contents as bytes.
    """
    qs = list(questions) if questions else STANDARD_QUESTIONS
    answers = data.get("answers", {})

    meta = {
        "candidate":               data.get("candidate_name", ""),
        "position":                data.get("position", ""),
        "referee":                 data.get("referee_name", ""),
        "referee_title":           data.get("referee_title", ""),
        "referee_previous":        data.get("referee_previous", ""),
        "referee_current_title":   data.get("referee_current_title"),
        "referee_current_company": data.get("referee_current_company"),
        "referee_previous_title":  data.get("referee_previous_title"),
        "referee_previous_company": data.get("referee_previous_company"),
        "date":                    data.get("reference_date", ""),
        "completed_by":            data.get("completed_by", "Tate McClenaghan"),
    }

    doc = Document()
    _set_page_margins(doc, top=0.6, bottom=0.7, left=0.85, right=0.85)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    _build_banner(doc, "Reference Check")
    _build_details(doc, meta)
    _build_questions(doc, qs, answers)
    _build_signoff(doc, meta)
    _build_footer(doc)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
