"""CV PDF generator — v2.

Creates branded, redacted CV PDFs with an Infinitas cover page.
Uses pymupdf (fitz) for proper PDF text redaction — preserves layout.
Uses AI to identify what to remove beyond regex patterns.
"""

import io
import re
from pathlib import Path

import fitz  # pymupdf


COVER_DOCX_PATH = Path(__file__).parent.parent / "templates" / "cv-cover-template.docx"

# Regex patterns for PII
EMAIL_RE = re.compile(r"\S+@\S+\.\S+")
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{3,4}")
URL_RE = re.compile(r"https?://\S+|www\.\S+|linkedin\.com/\S*", re.IGNORECASE)
ADDRESS_RE = re.compile(
    r"\d+\s+[\w\s]+(?:street|st|road|rd|avenue|ave|drive|dr|lane|ln|"
    r"crescent|cres|place|pl|way|terrace|tce|close|cl)\b",
    re.IGNORECASE,
)

# Section headers that signal "references" section
REFERENCES_HEADERS = {
    "references", "referees", "reference", "referee",
    "references available upon request",
    "references available on request",
}


def _create_cover_page(candidate_name: str, client_name: str) -> bytes:
    """Generate a branded cover page PDF from the DOCX template.

    Fills 'PROFILE OF' and 'Prepared for' paragraphs, converts to PDF via MS Graph.
    """
    from docx import Document

    doc = Document(str(COVER_DOCX_PATH))

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("PROFILE OF"):
            # Append candidate name to the existing run
            for run in para.runs:
                if "PROFILE OF" in run.text:
                    run.text = f"PROFILE OF {candidate_name.upper()}"
                    break
        elif text.startswith("Prepared for"):
            for run in para.runs:
                if "Prepared for" in run.text:
                    run.text = f"Prepared for {client_name}"
                    break

    # Save DOCX to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.getvalue()

    # Convert to PDF via MS Graph
    from ui import convert_docx_to_pdf
    pdf_bytes = convert_docx_to_pdf(docx_bytes, "cover.docx")
    if pdf_bytes is None:
        raise RuntimeError("Cover page PDF conversion failed. Check MS Graph auth.")
    return pdf_bytes


def _redact_pdf(pdf_bytes: bytes) -> bytes:
    """Redact PII from a PDF using pymupdf. Preserves layout.

    Finds and whites-out: emails, phone numbers, URLs, addresses.
    Removes entire references/referees section.
    Also removes all link annotations.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for page in doc:
        text = page.get_text()
        lines = text.split("\n")

        # --- Regex-based redaction ---
        for pattern in [EMAIL_RE, URL_RE]:
            for match in pattern.finditer(text):
                found = match.group()
                rects = page.search_for(found)
                for rect in rects:
                    page.add_redact_annot(rect, fill=(1, 1, 1))

        # Phone numbers — only redact if 7+ digits
        for match in PHONE_RE.finditer(text):
            digits = re.sub(r"\D", "", match.group())
            if len(digits) >= 7:
                found = match.group().strip()
                if found:
                    rects = page.search_for(found)
                    for rect in rects:
                        page.add_redact_annot(rect, fill=(1, 1, 1))

        # Addresses
        for match in ADDRESS_RE.finditer(text):
            rects = page.search_for(match.group())
            for rect in rects:
                page.add_redact_annot(rect, fill=(1, 1, 1))

        # --- References section detection ---
        in_references = False
        for line in lines:
            stripped = line.strip().lower().rstrip(":")
            if stripped in REFERENCES_HEADERS:
                in_references = True

            if in_references and line.strip():
                rects = page.search_for(line.strip())
                for rect in rects:
                    page.add_redact_annot(rect, fill=(1, 1, 1))

        # --- Remove link annotations ---
        annots = page.annots()
        if annots:
            for annot in annots:
                if annot.type[0] == 2:  # Link annotation
                    page.delete_annot(annot)

        # Apply all redactions
        page.apply_redactions()

    # Save
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)
    return buf.getvalue()


def _redact_pdf_with_ai(pdf_bytes: bytes) -> bytes:
    """AI-enhanced PDF redaction. Uses regex first, then AI for harder cases.

    AI identifies additional PII that regex misses (unusual phone formats,
    postal codes, suburb names used as addresses, etc.).
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    full_text = "\n".join(page.get_text() for page in doc)
    doc.close()

    # Get AI to identify what to remove
    from ai import redact_cv_text
    redacted_text = redact_cv_text(full_text)

    # Find lines that were removed by AI
    original_lines = [l.strip() for l in full_text.split("\n") if l.strip()]
    redacted_lines = set(l.strip() for l in redacted_text.split("\n") if l.strip())
    removed_lines = [l for l in original_lines if l not in redacted_lines]

    # Do regex redaction first
    partially_redacted = _redact_pdf(pdf_bytes)

    # Then redact AI-identified lines
    if removed_lines:
        doc = fitz.open(stream=partially_redacted, filetype="pdf")
        for page in doc:
            for line in removed_lines:
                if len(line) < 3:
                    continue
                rects = page.search_for(line)
                for rect in rects:
                    page.add_redact_annot(rect, fill=(1, 1, 1))
            page.apply_redactions()

        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        buf.seek(0)
        return buf.getvalue()

    return partially_redacted


def _redact_docx_regex(docx_bytes: bytes) -> bytes:
    """Redact a DOCX file using regex only (fast). Preserves formatting."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    in_references = False

    for para in doc.paragraphs:
        text_lower = para.text.strip().lower().rstrip(":")
        if text_lower in REFERENCES_HEADERS:
            in_references = True
        if in_references:
            for run in para.runs:
                run.text = ""
            continue

        for run in para.runs:
            original = run.text
            cleaned = EMAIL_RE.sub("", original)
            cleaned = URL_RE.sub("", cleaned)
            for match in PHONE_RE.finditer(cleaned):
                digits = re.sub(r"\D", "", match.group())
                if len(digits) >= 7:
                    cleaned = cleaned.replace(match.group(), "")
            cleaned = ADDRESS_RE.sub("", cleaned)
            cleaned = re.sub(r"  +", " ", cleaned).strip()
            if cleaned != original:
                run.text = cleaned

    for rel in list(doc.part.rels.values()):
        if "hyperlink" in str(rel.reltype).lower():
            try:
                doc.part.rels.pop(rel.rId)
            except Exception:
                pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _redact_docx_with_ai(docx_bytes: bytes) -> bytes:
    """Redact a DOCX file using AI + regex. Preserves formatting."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))

    # Get full text and AI-redacted version
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    from ai import redact_cv_text
    redacted_text = redact_cv_text(full_text)

    # Find removed lines
    original_lines = set(l.strip() for l in full_text.split("\n") if l.strip())
    redacted_lines = set(l.strip() for l in redacted_text.split("\n") if l.strip())
    removed_lines = original_lines - redacted_lines

    in_removed_section = False

    for para in doc.paragraphs:
        para_text = para.text.strip()

        # Check if this paragraph was removed by AI
        if para_text and para_text in removed_lines:
            for run in para.runs:
                run.text = ""
            in_removed_section = True
            continue

        if para_text and para_text in redacted_lines:
            in_removed_section = False

        if in_removed_section and para_text and para_text not in redacted_lines:
            for run in para.runs:
                run.text = ""
            continue

        # Regex PII stripping on runs
        for run in para.runs:
            original = run.text
            cleaned = EMAIL_RE.sub("", original)
            cleaned = URL_RE.sub("", cleaned)
            for match in PHONE_RE.finditer(cleaned):
                digits = re.sub(r"\D", "", match.group())
                if len(digits) >= 7:
                    cleaned = cleaned.replace(match.group(), "")
            cleaned = re.sub(r"  +", " ", cleaned).strip()
            if cleaned != original:
                run.text = cleaned

    # Remove hyperlinks
    for rel in list(doc.part.rels.values()):
        if "hyperlink" in str(rel.reltype).lower():
            try:
                doc.part.rels.pop(rel.rId)
            except Exception:
                pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _merge_pdfs(*pdf_bytes_list: bytes) -> bytes:
    """Merge multiple PDFs into one."""
    import pypdf
    writer = pypdf.PdfWriter()
    for pdf_bytes in pdf_bytes_list:
        for page in pypdf.PdfReader(io.BytesIO(pdf_bytes)).pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_cv_pdf(
    candidate_name: str,
    client_name: str,
    cv_file_bytes: bytes,
    cv_filename: str,
    use_ai_redaction: bool = False,
) -> bytes:
    """Generate a branded, redacted CV PDF.

    For PDF: pymupdf redaction (preserves original layout).
    For DOCX: regex (+ optional AI) redaction, convert to PDF via MS Graph.
    Both get a branded cover page prepended.

    Args:
        use_ai_redaction: If True, uses Claude API for deeper PII detection
            (slower ~10s per CV). If False, regex-only (instant).
    """
    cover_pdf = _create_cover_page(candidate_name, client_name)
    is_docx = cv_filename.lower().endswith(".docx")

    if is_docx:
        if use_ai_redaction:
            redacted_docx = _redact_docx_with_ai(cv_file_bytes)
        else:
            redacted_docx = _redact_docx_regex(cv_file_bytes)
        from ui import convert_docx_to_pdf
        cv_pdf = convert_docx_to_pdf(redacted_docx, cv_filename)
        if cv_pdf is None:
            raise RuntimeError("DOCX to PDF conversion failed. Check MS Graph auth.")
    else:
        if use_ai_redaction:
            cv_pdf = _redact_pdf_with_ai(cv_file_bytes)
        else:
            cv_pdf = _redact_pdf(cv_file_bytes)

    return _merge_pdfs(cover_pdf, cv_pdf)
